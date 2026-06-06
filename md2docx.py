"""
Markdown → Word 转换脚本
将详细复习资料 .md 转为格式化的 .docx
"""
import sys, os, re
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def create_word_from_md(md_path: str, docx_path: str):
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # -- Define styles --
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.35
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Heading styles
    for level, (size, bold, color) in enumerate([
        (Pt(22), True, RGBColor(0x1A, 0x3C, 0x6E)),  # Heading 1
        (Pt(16), True, RGBColor(0x2B, 0x57, 0x9A)),  # Heading 2
        (Pt(13), True, RGBColor(0x33, 0x66, 0x99)),  # Heading 3
        (Pt(11), True, RGBColor(0x44, 0x77, 0xAA)),  # Heading 4
    ], start=1):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = '微软雅黑'
        h_style.font.size = size
        h_style.font.bold = bold
        h_style.font.color.rgb = color
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        h_style.paragraph_format.space_before = Pt(12 if level > 1 else 18)
        h_style.paragraph_format.space_after = Pt(6)

    # Code style
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.0

    # -- Read and parse markdown --
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    in_mermaid_block = False
    in_table = False
    table_rows = []
    table_aligns = []

    while i < len(lines):
        line = lines[i].rstrip()

        # --- Mermaid blocks -> flowchart placeholder ---
        if line.startswith('```mermaid'):
            in_mermaid_block = True
            # Add a flowchart indicator paragraph
            p = doc.add_paragraph()
            run = p.add_run('[ 流程图 ]')
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(6)
            i += 1
            continue

        if in_mermaid_block:
            if line.startswith('```'):
                in_mermaid_block = False
                # End of mermaid block
                p = doc.add_paragraph()
                run = p.add_run('(请在 VS Code 中打开 .md 文件查看完整流程图，或安装 Markdown Preview Mermaid 插件)')
                run.italic = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                p.paragraph_format.left_indent = Cm(0.5)
                doc.add_paragraph()
                i += 1
                continue
            # Add mermaid content as styled code text
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
            i += 1
            continue

        # --- Code blocks ---
        if line.startswith('```') and not in_mermaid_block:
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='CodeBlock')
            i += 1
            continue

        # --- Tables ---
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]

            if i + 1 < len(lines) and '|---' in lines[i + 1]:
                # Header row + separator
                table_rows = [cells]
                # Parse alignment from separator
                sep_cells = [c.strip() for c in lines[i + 1].split('|')[1:-1]]
                table_aligns = []
                for sc in sep_cells:
                    if sc.startswith(':') and sc.endswith(':'):
                        table_aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
                    elif sc.endswith(':'):
                        table_aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
                    else:
                        table_aligns.append(WD_ALIGN_PARAGRAPH.LEFT)
                in_table = True
                i += 2
                continue
            elif in_table:
                table_rows.append(cells)
                i += 1
                continue
        elif in_table and line.strip() == '':
            # End table, render it
            _render_table(doc, table_rows, table_aligns)
            in_table = False
            table_rows = []
            table_aligns = []
            i += 1
            continue
        elif in_table:
            # End of table (non-empty, non-table line)
            _render_table(doc, table_rows, table_aligns)
            in_table = False
            table_rows = []
            table_aligns = []
            # fall through to process this line

        # --- Headings ---
        if line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # --- Blockquote ---
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            text = line[2:]
            _add_formatted_text(p, text, italic=True, color=RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue

        # --- Empty line ---
        if line.strip() == '':
            i += 1
            continue

        # --- Normal paragraph with inline formatting ---
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    # Render any remaining table
    if in_table and table_rows:
        _render_table(doc, table_rows, table_aligns)

    doc.save(docx_path)
    print(f'Generated: {docx_path}')


def _render_table(doc, rows, aligns):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols, style='Light Grid Accent 1')

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.rows[r_idx].cells[c_idx]
                # Clear default
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

                # Bold header row
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                if r_idx == 0:
                    run.bold = True
                    run.font.size = Pt(9.5)

                if aligns and c_idx < len(aligns):
                    p.alignment = aligns[c_idx]

    doc.add_paragraph()  # space after table


def _add_formatted_text(p, text, italic=False, color=None):
    """Add paragraph with inline bold (**text**) and code (`text`) support."""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = color


if __name__ == '__main__':
    md_file = sys.argv[1] if len(sys.argv) > 1 else '操作系统_详细复习资料.md'
    docx_file = md_file.replace('.md', '.docx')
    create_word_from_md(md_file, docx_file)
